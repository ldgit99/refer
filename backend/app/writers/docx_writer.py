"""DOCX writer with tracked-changes / annotated output (research.md §2.5.2).

Tracked mode wraps replacements in ``w:ins``/``w:del`` runs so Word shows them
under "Track Changes". Citation warnings are attached as real Word **comments**
(``w:comment`` in ``word/comments.xml``) so reviewers see them in the Review
pane; annotated mode additionally highlights the spot. Both modes are minimally
invasive: only targeted paragraphs are touched.
"""

from __future__ import annotations

import io

from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from lxml import etree

from app.writers.base import OutputMode, Patch

_AUTHOR = "refer"
_DATE = "2026-01-01T00:00:00Z"
_COMMENTS_PART_NAME = "/word/comments.xml"
_COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
_COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
_W_NS = nsmap["w"]


def _make_run_element(text: str):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _make_highlight_run(text: str, color: str = "yellow"):
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), color)
    rpr.append(hl)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _wrap_ins(run_el, rev_id: int):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), _AUTHOR)
    ins.set(qn("w:date"), _DATE)
    ins.append(run_el)
    return ins


def _wrap_del(text: str, rev_id: int):
    delel = OxmlElement("w:del")
    delel.set(qn("w:id"), str(rev_id))
    delel.set(qn("w:author"), _AUTHOR)
    delel.set(qn("w:date"), _DATE)
    r = OxmlElement("w:r")
    dtext = OxmlElement("w:delText")
    dtext.set(qn("xml:space"), "preserve")
    dtext.text = text
    r.append(dtext)
    delel.append(r)
    return delel


class _CommentsPart:
    """Accumulates Word comments and writes a ``word/comments.xml`` part.

    python-docx has no first-class comments API, so we build the comment
    elements in memory and, on :meth:`finalize`, serialise them into a new
    package part related from the document part.
    """

    def __init__(self, doc: Document) -> None:
        self._doc = doc
        self._root = OxmlElement("w:comments")
        self._next_id = 0

    def add(self, text: str) -> int:
        cid = self._next_id
        self._next_id += 1
        comment = OxmlElement("w:comment")
        comment.set(qn("w:id"), str(cid))
        comment.set(qn("w:author"), _AUTHOR)
        comment.set(qn("w:date"), _DATE)
        comment.set(qn("w:initials"), "RF")
        para = OxmlElement("w:p")
        para.append(_make_run_element(text))
        comment.append(para)
        self._root.append(comment)
        return cid

    def finalize(self) -> None:
        if self._next_id == 0:
            return
        document_part = self._doc.part
        blob = etree.tostring(
            self._root, xml_declaration=True, encoding="UTF-8", standalone=True
        )
        part = Part(
            PackURI(_COMMENTS_PART_NAME),
            _COMMENTS_CONTENT_TYPE,
            blob,
            document_part.package,
        )
        document_part.relate_to(part, _COMMENTS_REL_TYPE)


class DocxWriter:
    def apply(
        self,
        data: bytes,
        patches: list[Patch],
        mode: OutputMode = "tracked",
    ) -> bytes:
        doc = Document(io.BytesIO(data))
        paragraphs = doc.paragraphs
        rev = 1000
        comments = _CommentsPart(doc)

        for patch in patches:
            idx = patch.target.paragraph_index
            if idx < 0 or idx >= len(paragraphs):
                continue
            para = paragraphs[idx]

            if patch.kind in {"reference_replace", "doi_insert"}:
                rev = self._replace_paragraph(para, patch, mode, rev)
            elif patch.kind == "citation_comment":
                rev = self._annotate_paragraph(para, patch, mode, rev, comments)

        comments.finalize()
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _clear_runs(self, para) -> None:
        for r in list(para.runs):
            r._element.getparent().remove(r._element)

    def _replace_paragraph(self, para, patch: Patch, mode: OutputMode, rev: int) -> int:
        para_text = para.text
        original = patch.before or para_text
        new_text = patch.after
        if patch.before and patch.before not in para_text:
            return rev

        if mode == "final":
            self._clear_runs(para)
            para.add_run(para_text.replace(patch.before, new_text, 1) if patch.before else new_text)
            return rev

        p = para._p
        self._clear_runs(para)
        if mode == "tracked":
            p.append(_wrap_del(original, rev))
            rev += 1
            p.append(_wrap_ins(_make_run_element(new_text), rev))
            rev += 1
        else:  # annotated: keep original, append suggestion highlighted
            p.append(_make_run_element(original + "  "))
            p.append(_make_highlight_run(f"[제안: {new_text}]"))
        return rev

    def _annotate_paragraph(
        self, para, patch: Patch, mode: OutputMode, rev: int, comments: _CommentsPart
    ) -> int:
        note = patch.comment or patch.after or "검토 필요"
        p = para._p

        # Attach a real Word comment spanning the paragraph's run content.
        comment_id = comments.add(note)
        start = OxmlElement("w:commentRangeStart")
        start.set(qn("w:id"), str(comment_id))
        end = OxmlElement("w:commentRangeEnd")
        end.set(qn("w:id"), str(comment_id))
        ref_run = OxmlElement("w:r")
        ref = OxmlElement("w:commentReference")
        ref.set(qn("w:id"), str(comment_id))
        ref_run.append(ref)

        # commentRangeStart before existing content, end + reference after it.
        p.insert(0, start)
        p.append(end)
        p.append(ref_run)

        # In annotated mode also leave a visible inline highlight as a fallback
        # for viewers that do not render the comment pane.
        if mode == "annotated":
            p.append(_make_highlight_run(f" [{note}]"))
        return rev
