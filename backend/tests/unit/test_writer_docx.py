import io

from docx import Document

from app.writers.base import ParagraphRef, Patch
from app.writers.docx_writer import DocxWriter


def _docx(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _text(data: bytes) -> str:
    return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)


def test_reference_replace_final_contains_new_text() -> None:
    data = _docx(["Intro", "References", "Old, A. (2000). Old title."])
    patch = Patch(
        id="p1",
        kind="reference_replace",
        target=ParagraphRef(paragraph_index=2),
        before="Old, A. (2000). Old title.",
        after="Old, A. (2000). New APA title. Journal, 1(1), 1-2.",
        source="F2",
    )
    out = DocxWriter().apply(data, [patch], mode="final")
    assert "New APA title" in _text(out)


def test_final_mode_replaces_text() -> None:
    data = _docx(["References", "Old, A. (2000). Old."])
    patch = Patch(
        id="p1",
        kind="reference_replace",
        target=ParagraphRef(paragraph_index=1),
        before="Old, A. (2000). Old.",
        after="Brand new reference.",
        source="F2",
    )
    out = DocxWriter().apply(data, [patch], mode="final")
    assert "Brand new reference." in _text(out)


def test_annotated_mode_keeps_original() -> None:
    data = _docx(["References", "Old, A. (2000). Old."])
    patch = Patch(
        id="p1",
        kind="reference_replace",
        target=ParagraphRef(paragraph_index=1),
        before="Old, A. (2000). Old.",
        after="New ref.",
        source="F2",
    )
    out = DocxWriter().apply(data, [patch], mode="annotated")
    text = _text(out)
    assert "Old, A. (2000). Old." in text
    assert "New ref." in text


def test_no_patches_roundtrips() -> None:
    data = _docx(["A", "B"])
    out = DocxWriter().apply(data, [], mode="tracked")
    assert _text(out) == "A\nB"


def test_patch_with_missing_before_text_is_skipped() -> None:
    data = _docx(["References", "Actual paragraph."])
    patch = Patch(
        id="p1",
        kind="reference_replace",
        target=ParagraphRef(paragraph_index=1),
        before="Different paragraph.",
        after="Replacement that should not apply.",
        source="F2",
    )
    out = DocxWriter().apply(data, [patch], mode="final")
    text = _text(out)
    assert "Actual paragraph." in text
    assert "Replacement that should not apply." not in text
